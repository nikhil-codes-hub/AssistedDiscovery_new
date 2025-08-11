#!/usr/bin/env python3
"""
Convert USER_GUIDE.md to Word document
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import re
import os

def create_word_document():
    # Read the markdown file
    with open('USER_GUIDE.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create a new document
    doc = Document()
    
    # Set up styles
    setup_styles(doc)
    
    # Split content into lines
    lines = content.split('\n')
    
    # Process each line
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:  # Empty line
            i += 1
            continue
            
        if line.startswith('# '):  # Main title
            title = line[2:].strip()
            p = doc.add_heading(title, level=1)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
        elif line.startswith('## '):  # Section heading
            heading = line[3:].strip()
            doc.add_heading(heading, level=2)
            
        elif line.startswith('### '):  # Subsection heading
            heading = line[4:].strip()
            doc.add_heading(heading, level=3)
            
        elif line.startswith('#### '):  # Sub-subsection heading
            heading = line[5:].strip()
            doc.add_heading(heading, level=4)
            
        elif line.startswith('- '):  # Bullet point
            bullet_text = line[2:].strip()
            # Remove markdown formatting
            bullet_text = clean_markdown_text(bullet_text)
            p = doc.add_paragraph(bullet_text, style='List Bullet')
            
        elif re.match(r'^\d+\.', line):  # Numbered list
            list_text = re.sub(r'^\d+\.\s*', '', line)
            list_text = clean_markdown_text(list_text)
            p = doc.add_paragraph(list_text, style='List Number')
            
        elif line.startswith('```'):  # Code block
            # Find the end of code block
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph(code_text)
            p.style = 'Code'
            
        elif line.startswith('**Q:') or line.startswith('**Issue:'):  # FAQ or troubleshooting
            question = clean_markdown_text(line)
            p = doc.add_paragraph(question)
            p.style = 'Intense Quote'
            
        elif line.startswith('A:') or line.startswith('- **Solution**:') or line.startswith('- **Cause**:'):
            answer = clean_markdown_text(line)
            p = doc.add_paragraph(answer, style='Quote')
            
        elif line.startswith('---'):  # Horizontal rule - skip
            pass
            
        else:  # Regular paragraph
            if line:  # Only add non-empty lines
                text = clean_markdown_text(line)
                if text:
                    doc.add_paragraph(text)
        
        i += 1
    
    # Save the document
    output_path = 'AssistedDiscovery_USER_GUIDE.docx'
    doc.save(output_path)
    print(f"Word document created: {output_path}")
    return output_path

def setup_styles(doc):
    """Set up custom styles for the document"""
    
    # Code style
    try:
        code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_font = code_style.font
        code_font.name = 'Courier New'
        code_font.size = Pt(9)
        code_style.paragraph_format.left_indent = Inches(0.5)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
    except:
        pass  # Style might already exist

def clean_markdown_text(text):
    """Remove markdown formatting from text"""
    # Remove bold markers
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    
    # Remove italic markers
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    
    # Remove code markers
    text = re.sub(r'`(.*?)`', r'\1', text)
    
    # Remove link formatting
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    
    # Remove emojis (basic cleanup)
    text = re.sub(r'[🔍✅💾📊🏢📚🎯🤖🎨]', '', text)
    
    # Clean up extra spaces
    text = re.sub(r'\s+', ' ', text).strip()
    
    return text

if __name__ == "__main__":
    create_word_document()